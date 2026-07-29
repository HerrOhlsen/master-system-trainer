"""Build the printable black and white cheat sheet from the JSON data.

Layout: one dense A4 page. Header with the rules, the consonant table 0-9 with
its mnemonics, then all 100 words in five columns of twenty rows, so that every
column holds exactly two of the ten-blocks the book recommends learning in.
"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
import mastersystem  # noqa: E402

# Next to the repository by default, so the sheet lands beside the source book.
# Pass a different path as the first argument to override.
ROOT = Path(__file__).resolve().parent.parent
OUTPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT.parent / "master-system-spickzettel.docx"

FONT = "Calibri"
SHADE_BLOCK = "EAEAEA"
SHADE_HEAD = "D9D9D9"


def shade(cell, hex_color):
    element = OxmlElement("w:shd")
    element.set(qn("w:val"), "clear")
    element.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(element)


def set_borders(table, size=4, color="808080"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def tighten(paragraph, space_after=0, space_before=0):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.line_spacing = 1.0


def write(cell, runs, align=WD_ALIGN_PARAGRAPH.LEFT):
    """runs: list of (text, size_pt, bold, italic)."""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    tighten(paragraph)
    for text, size, bold, italic in runs:
        run = paragraph.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
    cell.vertical_alignment = 1  # center


def build():
    data = mastersystem.load()
    problems = mastersystem.check(data)
    if problems:
        print("Abbruch, Daten sind nicht sauber:")
        for problem in problems:
            print("  " + problem)
        return 1

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.1)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9)

    # --- Title -------------------------------------------------------------
    title = document.add_paragraph()
    tighten(title, space_after=1)
    run = title.add_run("Master-System")
    run.font.name = FONT
    run.font.size = Pt(17)
    run.bold = True
    run = title.add_run("   Merkwörter für die Zahlen 0 bis 99")
    run.font.name = FONT
    run.font.size = Pt(10.5)

    subtitle = document.add_paragraph()
    tighten(subtitle, space_after=5)
    run = subtitle.add_run("nach Christiane Stenger, Warum fällt das Schaf vom Baum?")
    run.font.name = FONT
    run.font.size = Pt(7.5)
    run.italic = True

    # --- Consonant table ---------------------------------------------------
    consonants = data["konsonanten"]
    table = document.add_table(rows=3, cols=10)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table, size=4)
    for index, item in enumerate(consonants):
        column_width = Cm(1.85)
        digit_cell = table.cell(0, index)
        letter_cell = table.cell(1, index)
        word_cell = table.cell(2, index)
        for cell in (digit_cell, letter_cell, word_cell):
            cell.width = column_width
        shade(digit_cell, SHADE_HEAD)
        write(digit_cell, [(str(item["ziffer"]), 15, True, False)], WD_ALIGN_PARAGRAPH.CENTER)
        write(letter_cell, [(item["anzeige"], 10.5, True, False)], WD_ALIGN_PARAGRAPH.CENTER)
        write(word_cell, [(item["merkwort"], 9, False, True)], WD_ALIGN_PARAGRAPH.CENTER)

    # --- Mnemonics and rules ----------------------------------------------
    mnemonics = document.add_paragraph()
    tighten(mnemonics, space_before=5, space_after=2)
    for index, item in enumerate(consonants):
        run = mnemonics.add_run(f"{item['ziffer']} ")
        run.font.name = FONT
        run.font.size = Pt(7.5)
        run.bold = True
        text = item["eselsbruecke"]
        separator = "   " if index < len(consonants) - 1 else ""
        run = mnemonics.add_run(text + separator)
        run.font.name = FONT
        run.font.size = Pt(7.5)

    for rule in data["regeln"][:3]:
        paragraph = document.add_paragraph()
        tighten(paragraph, space_after=0)
        run = paragraph.add_run("▪ " + rule)
        run.font.name = FONT
        run.font.size = Pt(8)

    spacer = document.add_paragraph()
    tighten(spacer, space_after=4)
    spacer.add_run("")

    # --- Main table: five blocks of twenty rows -----------------------------
    entries = {entry["zahl"]: entry for entry in data["eintraege"]}
    rows = 20
    blocks = 5
    main = document.add_table(rows=rows, cols=blocks * 2)
    main.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(main, size=4)

    for row in range(rows):
        main.rows[row].height = Cm(0.82)
        main.rows[row].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for block in range(blocks):
            number = block * 20 + row
            entry = entries[number]
            number_cell = main.cell(row, block * 2)
            word_cell = main.cell(row, block * 2 + 1)
            number_cell.width = Cm(1.0)
            word_cell.width = Cm(2.7)

            # Shade the second ten-block of each column so the blocks stay visible.
            if row >= 10:
                shade(number_cell, SHADE_BLOCK)
                shade(word_cell, SHADE_BLOCK)

            # The book uses single digits for 0-9; 00-09 would be a different code.
            label = str(number) if number < 10 else f"{number}"
            write(number_cell, [(label, 12, True, False)], WD_ALIGN_PARAGRAPH.CENTER)

            runs = [(entry["wort"], 12, True, False)]
            if entry["alternativen"]:
                runs.append(("  " + " · ".join(entry["alternativen"]), 7.5, False, True))
            write(word_cell, runs)

    footer = document.add_paragraph()
    tighten(footer, space_before=6)
    run = footer.add_run(
        "Fett = Stengers Vorschlag, kursiv = Alternativen. Eigene Merkwörter mit Bleistift daneben schreiben. "
        "Übe in Zehnerblöcken und immer in beide Richtungen: Zahl zu Bild und Bild zu Zahl. "
        "Für Zahlenreihen, die mit einer Null beginnen, lohnen sich später eigene Wörter für 00 bis 09."
    )
    run.font.name = FONT
    run.font.size = Pt(8)
    run.italic = True

    document.save(OUTPUT)
    print(f"Geschrieben: {OUTPUT}")
    return 0


if __name__ == "__main__":
    raise SystemExit(build())
